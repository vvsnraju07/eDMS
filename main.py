import os
import fitz  # PyMuPDF
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from a2wsgi import ASGIMiddleware
from docx import Document  

# Initialize FastAPI app
app = FastAPI()
origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
wsgi_app = ASGIMiddleware(app)

class SearchRequest(BaseModel):
    folder_name: str
    keyword: str

def load_pdfs(folder_path):
    documents = []
    for dirpath, _, filenames in os.walk(folder_path):
        for filename in filenames:
            if filename.endswith('.pdf') and not filename.endswith('_highlighted.pdf'):
                file_path = os.path.join(dirpath, filename)
                try:
                    doc = fitz.open(file_path)
                    full_text = ""
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        text = page.get_text("text")
                        full_text += text
                    documents.append({"filename": filename, "filepath": file_path, "content": full_text})
                    doc.close()
                except Exception as e:
                    print(f'Error processing {filename}: {str(e)}')
    return documents

def load_docx_files(folder_path):
    documents = []
    for dirpath, _, filenames in os.walk(folder_path):
        for filename in filenames:
            if filename.endswith('.docx'):
                file_path = os.path.join(dirpath, filename)
                try:
                    doc = Document(file_path)
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    documents.append({"filename": filename, "filepath": file_path, "content": '\n'.join(full_text)})
                except Exception as e:
                    print(f'Error processing {filename}: {str(e)}')
    return documents

def count_keyword_occurrences(documents, keyword):
    keyword_counts = {}
    keyword_contexts = {}
    for document in documents:
        text = document["content"]
        count = 0
        contexts = []
        start = 0
        while (index := text.lower().find(keyword.lower(), start)) != -1:
            count += 1
            start = index + len(keyword)
            # Extract context
            tokens = text.split()
            word_index = text[:index].count(' ')
            start_context = max(word_index - 5, 0)
            end_context = min(word_index + 11, len(tokens))
            context = ' '.join(tokens[start_context:end_context])
            contexts.append(context)
        if count > 0:
            keyword_counts[document["filename"]] = count
            keyword_contexts[document["filename"]] = contexts
    return keyword_counts, keyword_contexts

def sort_files_by_frequency(keyword_counts):
    sorted_files = sorted(keyword_counts.items(), key=lambda item: item[1], reverse=True)
    return sorted_files

def highlight_keyword_in_pdf(file_path, keyword):
    base, ext = os.path.splitext(file_path)
    highlighted_file_path = f"{base}_{keyword}_highlighted{ext}"

    if os.path.exists(highlighted_file_path):
        return highlighted_file_path

    doc = fitz.open(file_path)
    keyword_lower = keyword.lower()
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text_instances = page.search_for(keyword)
        if text_instances:
            print(f"Found {len(text_instances)} instances of keyword '{keyword}' on page {page_num + 1}")
        for inst in text_instances:
            highlight = page.add_highlight_annot(inst)
            highlight.update()

    doc.save(highlighted_file_path, garbage=4, deflate=True)
    doc.close()
    return highlighted_file_path

def highlight_keyword_in_word(file_path, keyword):
    base, ext = os.path.splitext(file_path)
    highlighted_file_path = f"{base}_{keyword}_highlighted{ext}"

    if os.path.exists(highlighted_file_path):
        return highlighted_file_path

    doc = Document(file_path)
    keyword_lower = keyword.lower()

    for para in doc.paragraphs:
        if keyword_lower in para.text.lower():
            inline = para.runs
            for i in range(len(inline)):
                if keyword_lower in inline[i].text.lower():
                    # Highlight text by modifying XML directly
                    inline[i]._element.get_or_add_rPr().get_or_add_highlight().set('val', 'yellow')

    doc.save(highlighted_file_path)
    return highlighted_file_path

@app.post('/search')
async def search(search_request: SearchRequest, request: Request):
    folder_name = search_request.folder_name
    keyword = search_request.keyword
    with open("path.txt", 'r') as file:
        path = file.readline().strip()
    folder_path = os.path.join(path, folder_name)

    if not os.path.exists(folder_path):
        raise HTTPException(status_code=404, detail="Folder not found")

    pdf_documents = load_pdfs(folder_path)
    docx_documents = load_docx_files(folder_path)

    documents = pdf_documents + docx_documents

    keyword_counts, keyword_contexts = count_keyword_occurrences(documents, keyword)
    sorted_files = sort_files_by_frequency(keyword_counts)

    result = []
    base_url = str(request.base_url)
    for filename, count in sorted_files:
        for doc in documents:
            if doc["filename"] == filename:
                if filename.endswith('.pdf'):
                    highlighted_file_path = highlight_keyword_in_pdf(doc["filepath"], keyword)
                    highlighted_url = f"{base_url}highlighted_pdfs/{folder_name}/{os.path.basename(highlighted_file_path)}"
                elif filename.endswith('.docx'):
                    highlighted_file_path = highlight_keyword_in_word(doc["filepath"], keyword)
                    highlighted_url = f"{base_url}highlighted_docx/{folder_name}/{os.path.basename(highlighted_file_path)}"

                result.append({
                    'filename': filename,
                    'filepath': doc["filepath"],
                    'count': count,
                    'highlighted_path': highlighted_url,
                    'contexts': keyword_contexts[filename]
                })

    return JSONResponse(content=result)


@app.get('/highlighted_pdfs/{folder_name}/{filename:path}')
async def open_highlighted_pdf_file(folder_name: str, filename: str):
    with open("path.txt", 'r') as file:
        path = file.readline().strip()
    folder_path = os.path.join(path, folder_name)

    file_path = os.path.join(folder_path, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(path=file_path, filename=filename, media_type='application/pdf', headers={"Content-Disposition": "inline"})


@app.get('/highlighted_docx/{folder_name}/{filename:path}')
async def open_highlighted_docx_file(folder_name: str, filename: str):
    with open("path.txt", 'r') as file:
        path = file.readline().strip()
    folder_path = os.path.join(path, folder_name)

    file_path = os.path.join(folder_path, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    # Serving the DOCX file with headers to try to open it in the browser
    headers = {
        "Content-Disposition": f"inline; filename={filename}"
    }
    return FileResponse(path=file_path, filename=filename,
                        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        headers=headers)

@app.get('/pdfs/{folder_name}/{filename:path}')
async def open_pdf_file(folder_name: str, filename: str):
    with open("path.txt", 'r') as file:
        path = file.readline().strip()
    folder_path = os.path.join(path, folder_name)

    file_path = os.path.join(folder_path, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(path=file_path, filename=filename, media_type='application/pdf', headers={"Content-Disposition": "inline"})


@app.get('/docx/{folder_name}/{filename:path}')
async def open_docx_file(folder_name: str, filename: str):
    with open("path.txt", 'r') as file:
        path = file.readline().strip()
    folder_path = os.path.join(path, folder_name)

    file_path = os.path.join(folder_path, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(path=file_path, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={"Content-Disposition": "inline"})


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
